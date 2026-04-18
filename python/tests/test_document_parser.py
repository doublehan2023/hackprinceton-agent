from __future__ import annotations

from docx import Document

from src.parsers.document_parser import parse_document, parse_text


def test_parse_text_creates_sections_from_numbered_headings() -> None:
    parsed = parse_text(
        "1. Confidentiality\n"
        "Confidential information must remain protected.\n\n"
        "2. Payment Terms\n"
        "Invoices are due within thirty days."
    )

    assert parsed.raw_text.startswith("1. Confidentiality")
    assert [section.title for section in parsed.sections] == ["1. Confidentiality", "2. Payment Terms"]
    assert [section.source_order for section in parsed.sections] == [1, 2]


def test_parse_text_falls_back_to_single_section_without_headings() -> None:
    parsed = parse_text(
        "Confidential information must remain protected.\n\n"
        "Invoices are due within thirty days of receipt."
    )

    assert len(parsed.sections) == 1
    assert parsed.sections[0].title is None
    assert "Invoices are due" in parsed.sections[0].body


def test_parse_text_avoids_false_positive_for_uppercase_clause_sentence() -> None:
    parsed = parse_text(
        "CONFIDENTIAL INFORMATION MUST REMAIN PROTECTED THROUGHOUT THE TERM OF THIS AGREEMENT\n"
        "The sponsor may disclose data only as required by law."
    )

    assert len(parsed.sections) == 1
    assert parsed.sections[0].title is None
    assert parsed.sections[0].body.startswith("CONFIDENTIAL INFORMATION")


def test_parse_docx_uses_heading_styles_for_sections(tmp_path) -> None:
    file_path = tmp_path / "contract.docx"
    document = Document()
    document.add_heading("Confidentiality", level=1)
    document.add_paragraph("Confidential information must remain protected.")
    document.add_heading("Payment Terms", level=1)
    document.add_paragraph("Invoices are due within thirty days.")
    document.save(file_path)

    parsed = parse_document(str(file_path), file_path.name)

    assert [section.title for section in parsed.sections] == ["Confidentiality", "Payment Terms"]
    assert [section.heading_level for section in parsed.sections] == [1, 1]
    assert parsed.sections[1].body == "Invoices are due within thirty days."
