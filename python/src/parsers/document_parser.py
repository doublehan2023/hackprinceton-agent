from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pydantic import BaseModel, Field

from src.pipeline.state import Section

_NUMERIC_HEADING_RE = re.compile(
    r"^(?:(?:section|article)\s+)?(\d+(?:\.\d+){0,2})(?:[\.\):])?\s+[A-Z][A-Za-z0-9/&,\-() ]{2,}$",
    re.IGNORECASE,
)
_PREFIX_HEADING_RE = re.compile(
    r"^(section|article)\s+([A-Za-z0-9]+)(?:[\.\):])?\s+[A-Z][A-Za-z0-9/&,\-() ]{2,}$",
    re.IGNORECASE,
)
_COLON_HEADING_RE = re.compile(r"^[A-Z][A-Za-z0-9/&,\-() ]{1,80}:$")


class ParsedDocument(BaseModel):
    raw_text: str
    sections: list[Section] = Field(default_factory=list)


def _normalize_text(text: str) -> str:
    return re.sub(r"\r\n?", "\n", text).strip()


def _compact_text(lines: list[str]) -> str:
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _finalize_sections(sections: list[Section], raw_text: str) -> list[Section]:
    cleaned = [
        Section(
            title=section.title,
            body=section.body.strip(),
            source_order=section.source_order,
            heading_level=section.heading_level,
        )
        for section in sections
        if section.body.strip()
    ]
    if cleaned:
        return cleaned
    normalized = _normalize_text(raw_text)
    if not normalized:
        return []
    return [Section(title=None, body=normalized, source_order=1)]


def _is_mostly_uppercase(text: str) -> bool:
    letters = [char for char in text if char.isalpha()]
    if len(letters) < 4:
        return False
    upper_ratio = sum(1 for char in letters if char.isupper()) / len(letters)
    return upper_ratio >= 0.8


def _looks_like_heading(text: str) -> tuple[bool, int | None]:
    candidate = text.strip()
    if not candidate or len(candidate) > 120:
        return False, None

    words = candidate.split()
    if len(words) > 12:
        return False, None
    if candidate.endswith(".") and not candidate.endswith("..."):
        return False, None

    numeric_match = _NUMERIC_HEADING_RE.match(candidate)
    if numeric_match:
        return True, numeric_match.group(1).count(".") + 1

    prefix_match = _PREFIX_HEADING_RE.match(candidate)
    if prefix_match:
        return True, 1

    if _COLON_HEADING_RE.match(candidate):
        return True, 1

    if _is_mostly_uppercase(candidate) and len(words) <= 6 and not re.search(r"[.!?]", candidate):
        return True, 1

    return False, None


def _sectionize_entries(entries: list[tuple[str, int | None, bool]]) -> list[Section]:
    sections: list[Section] = []
    current_title: str | None = None
    current_level: int | None = None
    current_body: list[str] = []

    def flush_section() -> None:
        nonlocal current_title, current_level, current_body
        body = _compact_text(current_body)
        if body:
            sections.append(
                Section(
                    title=current_title,
                    body=body,
                    source_order=len(sections) + 1,
                    heading_level=current_level,
                )
            )
        current_title = None
        current_level = None
        current_body = []

    for text, level, is_heading in entries:
        if is_heading:
            flush_section()
            current_title = text
            current_level = level
            continue
        if text == "":
            if current_body and current_body[-1] != "":
                current_body.append("")
            continue
        current_body.append(text)

    flush_section()
    return sections


def _structure_from_text(raw_text: str) -> list[Section]:
    normalized = _normalize_text(raw_text)
    if not normalized:
        return []

    entries: list[tuple[str, int | None, bool]] = []
    lines = normalized.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            entries.append(("", None, False))
            continue
        is_heading, level = _looks_like_heading(stripped)
        entries.append((stripped, level, is_heading))

    return _finalize_sections(_sectionize_entries(entries), normalized)


def _read_pdf(file_path: str) -> ParsedDocument:
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError("PyPDF2 is required to parse PDF files.") from exc

    reader = PdfReader(file_path)
    raw_text = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    return ParsedDocument(raw_text=raw_text, sections=_structure_from_text(raw_text))


def _docx_heading_level(paragraph) -> int | None:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    match = re.match(r"Heading\s+(\d+)", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def _read_docx(file_path: str) -> ParsedDocument:
    document = Document(file_path)
    raw_paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    raw_text = "\n".join(raw_paragraphs).strip()

    entries: list[tuple[str, int | None, bool]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            entries.append(("", None, False))
            continue

        heading_level = _docx_heading_level(paragraph)
        if heading_level is not None:
            entries.append((text, heading_level, True))
            continue

        is_heading, inferred_level = _looks_like_heading(text)
        entries.append((text, inferred_level, is_heading))

    return ParsedDocument(raw_text=raw_text, sections=_finalize_sections(_sectionize_entries(entries), raw_text))


def _read_txt(file_path: str) -> ParsedDocument:
    raw_text = Path(file_path).read_text(encoding="utf-8")
    return ParsedDocument(raw_text=raw_text, sections=_structure_from_text(raw_text))


def parse_text(text: str) -> ParsedDocument:
    normalized = _normalize_text(text)
    return ParsedDocument(raw_text=normalized, sections=_structure_from_text(normalized))


def parse_document(file_path: str, filename: str | None = None) -> ParsedDocument:
    suffix = (filename or Path(file_path).name).lower()
    if suffix.endswith(".pdf"):
        return _read_pdf(file_path)
    if suffix.endswith(".docx"):
        return _read_docx(file_path)
    if suffix.endswith(".txt"):
        return _read_txt(file_path)
    raise ValueError("Unsupported file format. Expected PDF, DOCX, or TXT.")


def extract_text(file_path: str, filename: str | None = None) -> str:
    return parse_document(file_path, filename).raw_text
